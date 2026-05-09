
# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ─── Couleurs SOGIS ───────────────────────────────────────────
VIOLET     = RGBColor(0x6B, 0x21, 0xA8)
BLEU       = RGBColor(0x0E, 0xA5, 0xE9)
GRIS_FOND  = RGBColor(0xF1, 0xF5, 0xF9)
BLANC      = RGBColor(0xFF, 0xFF, 0xFF)
GRIS_TEXTE = RGBColor(0x47, 0x55, 0x69)
ROUGE      = RGBColor(0xDC, 0x26, 0x26)
VERT       = RGBColor(0x16, 0xA3, 0x4A)
ORANGE     = RGBColor(0xEA, 0x58, 0x0C)

doc = Document()

# ─── Marges ───────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

# ─── Styles de base ───────────────────────────────────────────
style_normal = doc.styles['Normal']
style_normal.font.name = 'Calibri'
style_normal.font.size = Pt(11)
style_normal.font.color.rgb = GRIS_TEXTE

def set_para_spacing(para, before=6, after=6):
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after  = Pt(after)

def add_heading(text, level=1, color=VIOLET):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = color
    run.font.bold = True
    if level == 1:
        run.font.size = Pt(20)
        set_para_spacing(p, before=18, after=8)
    elif level == 2:
        run.font.size = Pt(15)
        set_para_spacing(p, before=14, after=6)
    elif level == 3:
        run.font.size = Pt(12)
        set_para_spacing(p, before=10, after=4)
    return p

def add_body(text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    set_para_spacing(p, before=3, after=3)
    return p

def add_bullet(text, level=0, marker='•'):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    p.paragraph_format.left_indent = Cm(0.75 + level * 0.5)
    set_para_spacing(p, before=2, after=2)
    return p

def add_numbered(text):
    p = doc.add_paragraph(style='List Number')
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    set_para_spacing(p, before=2, after=2)
    return p

def add_info_box(label, value, label_color=VIOLET):
    p = doc.add_paragraph()
    r1 = p.add_run(f'{label}: ')
    r1.font.bold = True
    r1.font.color.rgb = label_color
    r1.font.size = Pt(11)
    r2 = p.add_run(value)
    r2.font.size = Pt(11)
    set_para_spacing(p, before=2, after=2)
    return p

def add_divider():
    p = doc.add_paragraph('─' * 80)
    p.runs[0].font.color.rgb = GRIS_FOND
    p.runs[0].font.size = Pt(8)
    set_para_spacing(p, before=4, after=4)

def shade_cell(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def add_table(headers, rows, header_color='6B21A8'):
    n_cols = len(headers)
    t = doc.add_table(rows=1 + len(rows), cols=n_cols)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # En-tête
    hdr_row = t.rows[0]
    for i, h in enumerate(headers):
        c = hdr_row.cells[i]
        shade_cell(c, header_color)
        p = c.paragraphs[0]
        run = p.add_run(h)
        run.font.bold = True
        run.font.color.rgb = BLANC
        run.font.size = Pt(10)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Données
    for ri, row_data in enumerate(rows):
        row = t.rows[ri + 1]
        for ci, val in enumerate(row_data):
            c = row.cells[ci]
            if (ri % 2) == 0:
                shade_cell(c, 'F8FAFC')
            p = c.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(10)
    return t

# ═══════════════════════════════════════════════════════════════
# PAGE DE GARDE
# ═══════════════════════════════════════════════════════════════
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(p_title, before=60, after=10)
run = p_title.add_run('SOGIS SARL')
run.font.name = 'Calibri'
run.font.size = Pt(40)
run.font.bold = True
run.font.color.rgb = VIOLET

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p_sub.add_run('Solutions & Gestion Intégrée de Services')
run2.font.name = 'Calibri'
run2.font.size = Pt(16)
run2.font.color.rgb = BLEU
run2.font.italic = True
set_para_spacing(p_sub, before=4, after=40)

p_doc = doc.add_paragraph()
p_doc.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p_doc.add_run('GUIDE D\'UTILISATION & DOCUMENTATION TECHNIQUE')
run3.font.name = 'Calibri'
run3.font.size = Pt(14)
run3.font.bold = True
run3.font.color.rgb = GRIS_TEXTE
set_para_spacing(p_doc, before=0, after=6)

p_v = doc.add_paragraph()
p_v.alignment = WD_ALIGN_PARAGRAPH.CENTER
run4 = p_v.add_run(f'Version 1.0 — {datetime.datetime.now().strftime("%B %Y")}')
run4.font.size = Pt(11)
run4.font.color.rgb = GRIS_TEXTE
set_para_spacing(p_v, before=0, after=80)

p_conf = doc.add_paragraph()
p_conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
run5 = p_conf.add_run('Document confidentiel — Usage interne SOGIS SARL')
run5.font.size = Pt(9)
run5.font.italic = True
run5.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# TABLE DES MATIÈRES
# ═══════════════════════════════════════════════════════════════
add_heading('TABLE DES MATIÈRES', level=1, color=VIOLET)
sections_toc = [
    ('1.', 'Présentation de la plateforme SOGIS'),
    ('2.', 'Architecture technique'),
    ('3.', 'Accès et navigation pour les utilisateurs'),
    ('4.', 'Fonctionnalités utilisateurs'),
    ('   4.1', 'Formulaire de demande Business'),
    ('   4.2', 'Formulaire de demande Services'),
    ('   4.3', 'Suivi de dossier'),
    ('   4.4', 'Commentaires et avis'),
    ('5.', 'Panneau d\'administration'),
    ('   5.1', 'Connexion à l\'administration'),
    ('   5.2', 'Tableau de bord — Gestion des demandes'),
    ('   5.3', 'Tableau de bord — Gestion des commentaires'),
    ('   5.4', 'Export des données'),
    ('   5.5', 'Déconnexion'),
    ('6.', 'Configuration et maintenance (Supabase)'),
    ('   6.1', 'Accès au tableau de bord Supabase'),
    ('   6.2', 'Gestion des utilisateurs admin'),
    ('   6.3', 'Consultation des données'),
    ('   6.4', 'Politiques de sécurité (RLS)'),
    ('7.', 'Déploiement sur GitHub Pages'),
    ('8.', 'Glossaire & Informations de connexion'),
]
for num, titre in sections_toc:
    p = doc.add_paragraph()
    r1 = p.add_run(f'{num}  ')
    r1.font.bold = True
    r1.font.color.rgb = VIOLET
    r2 = p.add_run(titre)
    r2.font.color.rgb = GRIS_TEXTE
    set_para_spacing(p, before=2, after=2)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 1. PRÉSENTATION
# ═══════════════════════════════════════════════════════════════
add_heading('1. Présentation de la plateforme SOGIS', level=1)

add_body(
    'SOGIS SARL (Solutions & Gestion Intégrée de Services) est une entreprise camerounaise proposant '
    'deux pôles d\'activités complémentaires : le conseil et l\'accompagnement en affaires (Business & Finance) '
    'd\'une part, et la logistique événementielle et les services à la demande (Services & Events) d\'autre part.'
)
add_body(
    'La plateforme web SOGIS est l\'interface digitale officielle de l\'entreprise. Elle permet aux clients '
    'et prospects de soumettre leurs demandes en ligne, de suivre l\'avancement de leurs dossiers, '
    'de déposer des avis, et aux administrateurs de gérer l\'ensemble de ces interactions depuis '
    'un tableau de bord sécurisé.'
)

add_heading('Objectifs de la plateforme', level=2, color=BLEU)
add_bullet('Offrir un point d\'entrée unique pour toutes les demandes clients')
add_bullet('Assurer la traçabilité complète de chaque dossier via un ticket unique')
add_bullet('Permettre la modération des avis et témoignages clients')
add_bullet('Faciliter la gestion interne via un tableau de bord centralisé')
add_bullet('Garantir la sécurité des données grâce à Supabase et PostgreSQL')

add_heading('Pages et URL', level=2, color=BLEU)
add_table(
    ['Page', 'URL', 'Description'],
    [
        ['Accueil', '/sogis/', 'Page principale — présentation des deux pôles'],
        ['Business & Finance', '/sogis/business', 'Services financiers, juridiques et d\'investissement'],
        ['Services & Events', '/sogis/services', 'Logistique événementielle et services à la demande'],
        ['Suivi de dossier', '/sogis/tracking', 'Recherche d\'un dossier par numéro de ticket'],
        ['Administration', '/sogis/admin', 'Tableau de bord (accès sécurisé)'],
        ['Connexion admin', '/sogis/admin/login', 'Page de connexion administrateur'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 2. ARCHITECTURE TECHNIQUE
# ═══════════════════════════════════════════════════════════════
add_heading('2. Architecture technique', level=1)

add_body(
    'La plateforme SOGIS repose sur une architecture frontend/backend découplée, '
    'hébergée sur GitHub Pages pour le frontend et alimentée par Supabase en tant que '
    'backend-as-a-service.'
)

add_heading('Stack technologique', level=2, color=BLEU)
add_table(
    ['Couche', 'Technologie', 'Rôle'],
    [
        ['Frontend', 'React 19 + TypeScript', 'Interface utilisateur'],
        ['Build Tool', 'Vite 7', 'Compilation et bundling'],
        ['Styling', 'TailwindCSS 4', 'Design et responsive'],
        ['Animations', 'Framer Motion', 'Transitions et micro-animations'],
        ['Routing', 'React Router DOM 7', 'Navigation entre pages'],
        ['Backend', 'Supabase (PostgreSQL)', 'Base de données + Auth + API REST'],
        ['Authentification', 'Supabase Auth', 'Connexion admin sécurisée'],
        ['Hébergement', 'GitHub Pages', 'Déploiement frontend'],
        ['CI/CD', 'GitHub Actions', 'Déploiement automatique'],
        ['Internationalisation', 'LanguageContext (FR/EN)', 'Bilinguisme'],
    ]
)

add_heading('Base de données — Structure des tables', level=2, color=BLEU)

add_body('Table : requests (Demandes clients)', bold=True, color=VIOLET)
add_table(
    ['Champ', 'Type', 'Description'],
    [
        ['id', 'UUID', 'Identifiant unique auto-généré'],
        ['ticket_id', 'TEXT (unique)', 'Numéro de ticket transmis au client'],
        ['created_at', 'TIMESTAMPTZ', 'Date et heure de soumission'],
        ['name', 'TEXT', 'Nom complet du demandeur'],
        ['email', 'TEXT', 'Adresse email'],
        ['phone', 'TEXT', 'Numéro de téléphone (optionnel)'],
        ['service', 'TEXT', 'Service demandé'],
        ['message', 'TEXT', 'Description détaillée de la demande'],
        ['service_type', 'TEXT', '"business" ou "services"'],
        ['status', 'TEXT', 'pending / accepted / inprogress / completed'],
        ['status_history', 'JSONB', 'Historique des changements de statut'],
    ]
)

doc.add_paragraph()
add_body('Table : comments (Commentaires & Avis)', bold=True, color=VIOLET)
add_table(
    ['Champ', 'Type', 'Description'],
    [
        ['id', 'UUID', 'Identifiant unique auto-généré'],
        ['created_at', 'TIMESTAMPTZ', 'Date de soumission'],
        ['name', 'TEXT', 'Nom de l\'auteur'],
        ['email', 'TEXT', 'Email de l\'auteur'],
        ['rating', 'INT (1-5)', 'Note sur 5 étoiles'],
        ['comment', 'TEXT', 'Contenu du commentaire'],
        ['service_type', 'TEXT', '"business" ou "services"'],
        ['status', 'TEXT', 'pending / validated / rejected'],
    ]
)

add_heading('Sécurité des données (Row Level Security)', level=2, color=BLEU)
add_table(
    ['Profil', 'Table', 'Droits'],
    [
        ['Anonyme (public)', 'requests', 'INSERT et SELECT (toutes les demandes)'],
        ['Anonyme (public)', 'comments', 'INSERT + SELECT (uniquement "validated")'],
        ['Authentifié (admin)', 'requests', 'Tous droits (SELECT, INSERT, UPDATE, DELETE)'],
        ['Authentifié (admin)', 'comments', 'Tous droits (SELECT, INSERT, UPDATE, DELETE)'],
    ]
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 3. ACCÈS UTILISATEURS
# ═══════════════════════════════════════════════════════════════
add_heading('3. Accès et navigation pour les utilisateurs', level=1)

add_body(
    'La plateforme est publiquement accessible via l\'URL officielle hébergée sur GitHub Pages. '
    'Aucune création de compte n\'est requise pour soumettre une demande ou un commentaire.'
)

add_heading('Navigateur compatible', level=2, color=BLEU)
add_bullet('Google Chrome (recommandé)')
add_bullet('Mozilla Firefox')
add_bullet('Microsoft Edge')
add_bullet('Safari (iOS & macOS)')
add_body('La plateforme est entièrement responsive et accessible sur mobile.', italic=True)

add_heading('Langue d\'interface', level=2, color=BLEU)
add_body(
    'La plateforme supporte le Français et l\'Anglais. '
    'Le bouton de changement de langue (FR / EN) est visible en haut à droite de la barre de navigation.'
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 4. FONCTIONNALITÉS UTILISATEURS
# ═══════════════════════════════════════════════════════════════
add_heading('4. Fonctionnalités utilisateurs', level=1)

# 4.1
add_heading('4.1 Formulaire de demande Business', level=2, color=BLEU)
add_body('Accessible depuis la page "Business & Finance" → bouton "Faire une demande".')
add_body('Champs du formulaire :')
add_bullet('Nom complet')
add_bullet('Adresse email')
add_bullet('Numéro de téléphone')
add_bullet('Type de service Business (liste déroulante)')
add_bullet('Description de la demande (zone de texte libre)')
add_body('Après soumission :', bold=True)
add_bullet('Un numéro de ticket unique est généré (ex. : BIZ-2026-XXXX)')
add_bullet('Le ticket est affiché à l\'écran — à conserver pour le suivi')
add_bullet('La demande est enregistrée dans Supabase avec statut "En attente"')
add_bullet('L\'administrateur en est notifié via le tableau de bord')

# 4.2
add_heading('4.2 Formulaire de demande Services & Events', level=2, color=BLEU)
add_body('Accessible depuis la page "Services & Events" → bouton "Faire une demande".')
add_body('Fonctionnement identique au formulaire Business. Le ticket généré est préfixé SRV-.')

# 4.3
add_heading('4.3 Suivi de dossier', level=2, color=BLEU)
add_body('Accessible depuis le bouton "Suivre mon Dossier" dans la barre de navigation, ou via /sogis/tracking.')
add_body('Étapes :')
add_numbered('Saisir le numéro de ticket reçu lors de la soumission')
add_numbered('Cliquer sur "Rechercher"')
add_numbered('La fiche du dossier s\'affiche avec : nom, service, date, statut actuel et historique des statuts')
add_body('Statuts possibles :', bold=True)
add_table(
    ['Statut', 'Signification', 'Couleur'],
    [
        ['En attente (pending)', 'Demande reçue, pas encore traitée', '🟡 Jaune'],
        ['Accepté (accepted)', 'Demande prise en charge par SOGIS', '🔵 Bleu'],
        ['En cours (inprogress)', 'Traitement actif du dossier', '🟣 Violet'],
        ['Finalisé (completed)', 'Dossier clôturé et service rendu', '🟢 Vert'],
    ]
)

# 4.4
add_heading('4.4 Commentaires et avis clients', level=2, color=BLEU)
add_body('Un formulaire de commentaires est disponible sur les pages Business et Services.')
add_body('Champs :')
add_bullet('Nom')
add_bullet('Email')
add_bullet('Note (1 à 5 étoiles)')
add_bullet('Commentaire libre')
add_body(
    'Les commentaires soumis sont en attente de validation par l\'administrateur. '
    'Seuls les commentaires validés sont affichés publiquement sur le site.',
    italic=True, color=ORANGE
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 5. PANNEAU D'ADMINISTRATION
# ═══════════════════════════════════════════════════════════════
add_heading('5. Panneau d\'administration', level=1)
add_body(
    'Le panneau d\'administration est l\'espace centralisé permettant à l\'équipe SOGIS de gérer '
    'l\'ensemble des demandes clients et des commentaires reçus via la plateforme.'
)

# 5.1
add_heading('5.1 Connexion à l\'administration', level=2, color=BLEU)
add_body('URL d\'accès :', bold=True)
add_info_box('URL', 'https://nemgp.github.io/sogis/admin/login', label_color=BLEU)
add_body('Identifiants par défaut :', bold=True)
add_info_box('Email', 'admin@sogis.cm')
add_info_box('Mot de passe', 'Sogis@2026Admin', label_color=ROUGE)
add_body(
    '⚠️ Il est fortement recommandé de changer ce mot de passe après la première connexion. '
    'Voir section 6.2 pour la procédure.',
    color=ROUGE, italic=True
)
add_body('Procédure de connexion :', bold=True)
add_numbered('Accéder à l\'URL de connexion')
add_numbered('Saisir l\'email et le mot de passe administrateur')
add_numbered('Cliquer sur "Se connecter"')
add_numbered('En cas de succès, redirection automatique vers le tableau de bord')
add_numbered('En cas d\'échec, un message d\'erreur s\'affiche (email ou mot de passe incorrect)')
add_body(
    'Note : La session est persistante. Vous restez connecté jusqu\'à ce que vous cliquiez sur '
    '"Déconnexion" ou fermiez le navigateur.',
    italic=True
)

# 5.2
add_heading('5.2 Tableau de bord — Gestion des demandes', level=2, color=BLEU)
add_body(
    'L\'onglet "Demandes" est l\'onglet par défaut du tableau de bord. '
    'Il liste toutes les demandes clients reçues, avec leurs informations complètes.'
)

add_body('Filtres disponibles :', bold=True)
add_bullet('Par statut : Toutes / En attente / Accepté / En cours / Finalisé')
add_bullet('Par type : Toutes / Business / Services')

add_body('Informations affichées par demande :', bold=True)
add_bullet('Numéro de ticket')
add_bullet('Type (Business ou Services) — badge coloré')
add_bullet('Statut actuel — badge coloré')
add_bullet('Nom, email, téléphone du demandeur')
add_bullet('Date et heure de soumission')
add_bullet('Service demandé')
add_bullet('Message complet')

add_body('Actions disponibles :', bold=True)
add_table(
    ['Action', 'Condition', 'Effet'],
    [
        ['Accepter', 'Statut = En attente', 'Passe le statut à "Accepté"'],
        ['Démarrer', 'Statut = Accepté', 'Passe le statut à "En cours"'],
        ['Finaliser', 'Statut = En cours', 'Passe le statut à "Finalisé"'],
        ['Supprimer', 'Toujours disponible', 'Supprime définitivement la demande (confirmation requise)'],
    ]
)
add_body(
    'Note : Chaque changement de statut est enregistré dans l\'historique du dossier '
    '(accessible par le client depuis la page de suivi).',
    italic=True
)

# 5.3
add_heading('5.3 Tableau de bord — Gestion des commentaires', level=2, color=BLEU)
add_body(
    'L\'onglet "Commentaires" affiche tous les avis en attente de modération. '
    'Les commentaires validés sont publiés publiquement sur le site ; '
    'les commentaires rejetés sont archivés et masqués.'
)

add_body('Filtres disponibles :')
add_bullet('Tous / Business / Services')

add_body('Actions disponibles :', bold=True)
add_table(
    ['Action', 'Icône', 'Effet'],
    [
        ['Valider', '✅ Vert', 'Publie le commentaire sur le site public (statut → validated)'],
        ['Rejeter', '❌ Rouge', 'Archive le commentaire, non publié (statut → rejected)'],
        ['Supprimer', '🗑️ Corbeille', 'Supprime définitivement le commentaire (confirmation requise)'],
    ]
)

add_body('Historique :', bold=True)
add_body(
    'En bas de l\'onglet Commentaires, deux sections collapsibles permettent de consulter '
    'l\'historique des commentaires Validés et Rejetés. Chaque ligne peut être supprimée définitivement.'
)

# 5.4
add_heading('5.4 Export des données', level=2, color=BLEU)
add_body(
    'Un bouton "Export Excel (CSV)" est disponible dans l\'onglet Demandes. '
    'Il génère un fichier CSV contenant toutes les demandes actuellement affichées, '
    'avec les colonnes suivantes : Ticket ID, Date, Nom, Email, Téléphone, Service, Message, Type, Statut.'
)
add_body('Le fichier est nommé automatiquement : SOGIS_Demandes_YYYY-MM-DD.csv')
add_body(
    'Pour ouvrir le CSV dans Excel, utiliser "Données → Importer un fichier texte/CSV" '
    'et sélectionner le séparateur "Virgule" avec encodage UTF-8.',
    italic=True
)

# 5.5
add_heading('5.5 Déconnexion', level=2, color=BLEU)
add_body(
    'Le bouton "Déconnexion" est affiché en haut à droite du tableau de bord. '
    'Cliquer dessus met fin à la session et redirige vers la page de connexion.'
)
add_body(
    '⚠️ En cas d\'inactivité prolongée, la session expire automatiquement côté Supabase. '
    'Vous serez alors redirigé vers /admin/login à votre prochain accès à /admin.',
    italic=True, color=ORANGE
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 6. CONFIGURATION ET MAINTENANCE
# ═══════════════════════════════════════════════════════════════
add_heading('6. Configuration et maintenance (Supabase)', level=1)

add_body(
    'Supabase est le backend de la plateforme SOGIS. Il héberge la base de données PostgreSQL, '
    'gère l\'authentification des administrateurs, et expose une API REST automatique.'
)

add_heading('6.1 Accès au tableau de bord Supabase', level=2, color=BLEU)
add_info_box('URL Dashboard', 'https://supabase.com/dashboard/project/fcdezpuhpsrdeihhngvi', BLEU)
add_info_box('Compte', 'Connecté avec le compte Google de SOGIS (nemgp)', GRIS_TEXTE)
add_info_box('Project URL', 'https://fcdezpuhpsrdeihhngvi.supabase.co', BLEU)

add_heading('6.2 Gestion des utilisateurs administrateurs', level=2, color=BLEU)
add_body('Pour ajouter un nouvel administrateur :', bold=True)
add_numbered('Se connecter sur Supabase Dashboard')
add_numbered('Aller dans Authentication → Users')
add_numbered('Cliquer sur "Add user" → "Create new user"')
add_numbered('Renseigner l\'email et un mot de passe fort')
add_numbered('Activer "Auto Confirm User" et cliquer "Create user"')
add_body('Pour changer un mot de passe :', bold=True)
add_numbered('Aller dans Authentication → Users')
add_numbered('Cliquer sur l\'utilisateur')
add_numbered('Cliquer "Send password recovery" (email envoyé) ou modifier manuellement')

add_heading('6.3 Consultation des données via Supabase', level=2, color=BLEU)
add_body(
    'Il est possible de consulter et modifier directement les données dans Supabase :'
)
add_numbered('Aller dans Table Editor')
add_numbered('Sélectionner la table "requests" ou "comments"')
add_numbered('Les données sont affichées sous forme de tableau éditable')
add_body('Pour exécuter une requête SQL personnalisée :', bold=True)
add_numbered('Aller dans SQL Editor')
add_numbered('Écrire et exécuter la requête souhaitée')
add_body('Exemple — Lister les demandes en attente :', italic=True)
p_code = doc.add_paragraph('SELECT ticket_id, name, email, service, created_at FROM requests WHERE status = \'pending\' ORDER BY created_at DESC;')
p_code.runs[0].font.name = 'Courier New'
p_code.runs[0].font.size = Pt(9)
p_code.paragraph_format.left_indent = Cm(1)

add_heading('6.4 Politiques de sécurité (RLS)', level=2, color=BLEU)
add_body(
    'Row Level Security (RLS) est activé sur les deux tables. '
    'Les politiques définissent qui peut lire/écrire quelles données :'
)
add_bullet('Les utilisateurs non connectés (anon) peuvent soumettre des demandes et commentaires')
add_bullet('Les utilisateurs non connectés ne peuvent lire que les commentaires "validated"')
add_bullet('Les administrateurs authentifiés ont accès complet à toutes les données')
add_body(
    'Ces politiques sont gérées dans Supabase Dashboard → Authentication → Policies.',
    italic=True
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 7. DÉPLOIEMENT
# ═══════════════════════════════════════════════════════════════
add_heading('7. Déploiement sur GitHub Pages', level=1)

add_body(
    'La plateforme est déployée automatiquement sur GitHub Pages à chaque push sur la branche "main". '
    'Le processus est géré par GitHub Actions (.github/workflows/deploy.yml).'
)

add_heading('Processus de déploiement', level=2, color=BLEU)
add_numbered('Modifier le code source en local')
add_numbered('Commiter et pousser sur GitHub : git add . && git commit -m "..." && git push')
add_numbered('GitHub Actions déclenche automatiquement le build (npm run build)')
add_numbered('Le résultat est publié sur GitHub Pages en 1-3 minutes')
add_info_box('URL de production', 'https://nemgp.github.io/sogis/', BLEU)

add_heading('Développement local', level=2, color=BLEU)
add_numbered('Ouvrir un terminal dans le dossier du projet')
add_numbered('Exécuter : npm run dev')
add_numbered('Accéder à : http://localhost:5173/sogis/')
add_body('Pré-requis : Node.js 18+ et npm installés.', italic=True)

add_heading('Variables d\'environnement', level=2, color=BLEU)
add_body('Le fichier .env à la racine du projet contient :', bold=True)
p_code2 = doc.add_paragraph(
    'VITE_SUPABASE_URL=https://fcdezpuhpsrdeihhngvi.supabase.co\n'
    'VITE_SUPABASE_ANON_KEY=sb_publishable_rfKTT_l3XGzvVn3X76g7XQ_AB0PvYFB'
)
p_code2.runs[0].font.name = 'Courier New'
p_code2.runs[0].font.size = Pt(9)
p_code2.paragraph_format.left_indent = Cm(1)
add_body(
    '⚠️ Ces variables sont exposées côté client (préfixe VITE_). '
    'La clé anon est publique par design — la sécurité est assurée par les politiques RLS de Supabase.',
    italic=True, color=ORANGE
)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# 8. GLOSSAIRE & INFORMATIONS DE CONNEXION
# ═══════════════════════════════════════════════════════════════
add_heading('8. Glossaire & Informations de connexion', level=1)

add_heading('Glossaire', level=2, color=BLEU)
add_table(
    ['Terme', 'Définition'],
    [
        ['Ticket ID', 'Numéro unique identifiant une demande (ex. BIZ-2026-XXXX ou SRV-2026-XXXX)'],
        ['RLS', 'Row Level Security — politique de sécurité au niveau des lignes de la BDD'],
        ['Supabase', 'Plateforme backend open-source basée sur PostgreSQL'],
        ['Anon key', 'Clé publique Supabase utilisée pour les appels non authentifiés'],
        ['GitHub Pages', 'Service d\'hébergement statique gratuit de GitHub'],
        ['GitHub Actions', 'Outil CI/CD intégré à GitHub pour l\'automatisation'],
        ['React', 'Bibliothèque JavaScript pour construire des interfaces utilisateur'],
        ['TypeScript', 'Superset de JavaScript avec typage statique'],
        ['Vite', 'Outil de build moderne pour les applications web'],
        ['TailwindCSS', 'Framework CSS utilitaire'],
    ]
)

add_heading('Récapitulatif des accès', level=2, color=ROUGE)
add_table(
    ['Ressource', 'URL / Identifiant', 'Notes'],
    [
        ['Site public', 'https://nemgp.github.io/sogis/', 'Accessible à tous'],
        ['Admin login', 'https://nemgp.github.io/sogis/admin/login', 'Réservé SOGIS'],
        ['Email admin', 'admin@sogis.cm', '—'],
        ['Mot de passe', 'Sogis@2026Admin', '⚠️ À changer'],
        ['Supabase Dashboard', 'https://supabase.com/dashboard\n/project/fcdezpuhpsrdeihhngvi', 'Compte Google SOGIS'],
        ['Repo GitHub', 'https://github.com/nemgp/sogis', 'Code source'],
        ['Dev local', 'http://localhost:5173/sogis/', 'npm run dev requis'],
    ],
    header_color='DC2626'
)

# Pied de page
doc.add_paragraph()
add_divider()
p_footer = doc.add_paragraph()
p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_f = p_footer.add_run(
    f'SOGIS SARL — Guide d\'utilisation v1.0 — {datetime.datetime.now().strftime("%d/%m/%Y")} — Confidentiel'
)
run_f.font.size = Pt(9)
run_f.font.italic = True
run_f.font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)

# ─── Sauvegarde ───────────────────────────────────────────────
output_path = r'C:\Users\mngue\.gemini\antigravity\scratch\sogis\SOGIS_Guide_Utilisation.docx'
doc.save(output_path)
print(f'Document cree : {output_path}')
